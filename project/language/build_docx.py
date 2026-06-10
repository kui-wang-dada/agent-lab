"""把方言客服系统-技术方案.md 转成排版友好的 .docx。

绕开 pandoc 的字体 + ASCII 框图问题：
- 中文字体强制设为黑体（标题）/ 宋体（正文）
- ASCII 框图段落改为"流程卡片"样式（居中段落 + 箭头 + 灰色背景）
- 表格用 docx 原生 Light Grid Accent 样式
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).parent
SRC = ROOT / "方言客服系统-技术方案.md"
DST = ROOT / "方言客服系统-技术方案.docx"

# 中文字体策略：macOS 系统字体
FONT_HEADING = "Heiti SC"      # 标题用黑体
FONT_BODY = "Songti SC"        # 正文用宋体
FONT_MONO = "Menlo"            # 代码用 Menlo

# ASCII 框线字符 —— 识别用
BOX_CHARS = set("┌┐└┘├┤┬┴┼─│▼▲◄►←→↓↑")


def set_run_font(run, font_name: str, size_pt: float | None = None,
                 bold: bool = False, color: tuple[int, int, int] | None = None):
    """设置 run 的中英文字体 + 样式。"""
    run.font.name = font_name
    # 关键：中文字符必须单独设置 east-asia 字体，否则会 fallback 到默认
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)

    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph_with_runs(doc, text: str, font: str = FONT_BODY,
                            size: float = 11, bold: bool = False,
                            align=WD_ALIGN_PARAGRAPH.LEFT,
                            color: tuple[int, int, int] | None = None,
                            space_after: float = 4):
    """添加带统一字体的段落。"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_run_font(run, font, size, bold, color)
    return p


def add_inline_runs(p, text: str, base_font: str = FONT_BODY, base_size: float = 11):
    """解析 inline markdown（**bold**、`code`）添加到段落。"""
    # 匹配 **xxx** 或 `xxx`
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, base_font, base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, FONT_MONO, base_size - 0.5)
            # 浅灰背景
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F2F2F2")
            run._element.get_or_add_rPr().append(shd)
        else:
            run = p.add_run(part)
            set_run_font(run, base_font, base_size)


def set_cell_shading(cell, fill_color: str):
    """设置单元格背景色（hex without #）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def add_card_paragraph(doc, lines: list[str]):
    """ASCII 框图段落 → 流程卡片：居中、浅灰背景、单倍行距。"""
    # 提取真正有信息的文字行（剔除纯框线）
    info_lines = []
    for line in lines:
        # 剔除纯框线字符 + 空白
        clean = "".join(ch for ch in line if ch not in BOX_CHARS).strip()
        if clean:
            info_lines.append(clean)

    if not info_lines:
        return

    # 用一个单格表格做"卡片"，灰色背景
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F5F5F7")

    # 设单元格边框
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D0D0D0")
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # 清掉单元格默认段落
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    # 用 ▼ 表示流程方向
    for i, line in enumerate(info_lines):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(line)
        set_run_font(run, FONT_BODY, 10.5, bold=("：" in line and len(line) < 30))

        if i < len(info_lines) - 1:
            arrow_p = cell.add_paragraph()
            arrow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow_p.paragraph_format.space_after = Pt(2)
            arrow_run = arrow_p.add_run("▼")
            set_run_font(arrow_run, FONT_BODY, 9, color=(150, 150, 150))

    # 卡片后空一行
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_code_block(doc, lines: list[str]):
    """普通代码块：等宽字体 + 浅灰背景。"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F7F7F7")
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line if line else " ")
        set_run_font(run, FONT_MONO, 9.5)
    doc.add_paragraph()


def parse_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """解析 markdown 表格。"""
    rows = [
        [c.strip() for c in re.split(r"(?<!\\)\|", line.strip().strip("|"))]
        for line in table_lines
    ]
    header = rows[0]
    # 跳过分隔行（rows[1] 是 |---|---|）
    body = rows[2:] if len(rows) > 2 else []
    return header, body


def add_table(doc, header: list[str], body: list[list[str]]):
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表头
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline_runs(p, h, base_font=FONT_HEADING, base_size=10.5)
        # header 字体加粗
        for run in p.runs:
            run.bold = True
        set_cell_shading(cell, "E8E8EC")

    # 表体
    for r, row in enumerate(body):
        for c, val in enumerate(row):
            if c >= len(header):
                break
            cell = table.rows[1 + r].cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
            p = cell.add_paragraph()
            p.paragraph_format.line_spacing = 1.3
            # 单元格里可能有 <br> 换行
            for j, segment in enumerate(val.split("<br>")):
                if j > 0:
                    p = cell.add_paragraph()
                    p.paragraph_format.line_spacing = 1.3
                add_inline_runs(p, segment, base_font=FONT_BODY, base_size=10)

    doc.add_paragraph()


def add_blockquote(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    add_inline_runs(p, text, base_font=FONT_BODY, base_size=10.5)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    # 左边竖线
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:color"), "B0B0B0")
    left.set(qn("w:space"), "8")
    pBdr.append(left)
    pPr.append(pBdr)


def add_list_item(doc, text: str, level: int = 0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    bullet_run = p.add_run("• ")
    set_run_font(bullet_run, FONT_BODY, 11)
    add_inline_runs(p, text, base_font=FONT_BODY, base_size=11)


def add_heading(doc, text: str, level: int):
    sizes = {1: 20, 2: 16, 3: 13.5, 4: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, FONT_HEADING, sizes.get(level, 11), bold=True,
                 color=(0x1F, 0x2E, 0x4D) if level <= 2 else (0x2C, 0x3E, 0x66))


def convert():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()

    # 设置默认页面边距
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            # 判断是否 ASCII 框图
            joined = "\n".join(block)
            if any(ch in joined for ch in BOX_CHARS):
                add_card_paragraph(doc, block)
            else:
                add_code_block(doc, block)
            continue

        # 表格
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?[\s\-:|]+\|", lines[i + 1]):
            table_lines = []
            while i < len(lines) and "|" in lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            header, body = parse_table(table_lines)
            add_table(doc, header, body)
            continue

        # 标题
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped.lstrip("#").strip()
            add_heading(doc, title, level)
            i += 1
            continue

        # 分隔线
        if stripped == "---":
            i += 1
            # 加一条细灰线
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "CCCCCC")
            bottom.set(qn("w:space"), "1")
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # 引用块
        if stripped.startswith(">"):
            quoted = stripped.lstrip(">").strip()
            add_blockquote(doc, quoted)
            i += 1
            continue

        # 无序列表（- 开头）
        m = re.match(r"^(\s*)-\s+(.+)$", line)
        if m:
            indent = len(m.group(1))
            content = m.group(2)
            add_list_item(doc, content, level=indent // 2)
            i += 1
            continue

        # 普通段落
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline_runs(p, stripped, base_font=FONT_BODY, base_size=11)
        i += 1

    doc.save(DST)
    print(f"✓ Generated: {DST}")
    print(f"  Size: {DST.stat().st_size:,} bytes")


if __name__ == "__main__":
    convert()
